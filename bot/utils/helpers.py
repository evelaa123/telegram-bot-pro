"""
Helper utilities for the bot.
"""
import html
import re
from typing import Union, List
from aiogram.types import Message, BufferedInputFile


def format_number(num: Union[int, float], precision: int = 2) -> str:
    """
    Format number with thousand separators.
    
    Args:
        num: Number to format
        precision: Decimal precision for floats
        
    Returns:
        Formatted string
    """
    if isinstance(num, float):
        return f"{num:,.{precision}f}"
    return f"{num:,}"


def truncate_text(text: str, max_length: int = 100, suffix: str = "...") -> str:
    """
    Truncate text to max length.
    
    Args:
        text: Text to truncate
        max_length: Maximum length
        suffix: Suffix to add if truncated
        
    Returns:
        Truncated text
    """
    if len(text) <= max_length:
        return text
    return text[:max_length - len(suffix)] + suffix


def escape_html(text: str) -> str:
    """
    Escape HTML special characters for Telegram HTML parse mode.
    
    Args:
        text: Text to escape
        
    Returns:
        Escaped text
    """
    return html.escape(text)


def format_duration(seconds: int) -> str:
    """
    Format duration in human readable format.
    
    Args:
        seconds: Duration in seconds
        
    Returns:
        Formatted string like "2h 30m" or "45s"
    """
    if seconds < 60:
        return f"{seconds}s"
    elif seconds < 3600:
        minutes = seconds // 60
        secs = seconds % 60
        return f"{minutes}m {secs}s" if secs else f"{minutes}m"
    else:
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        return f"{hours}h {minutes}m" if minutes else f"{hours}h"


def format_file_size(bytes_size: int) -> str:
    """
    Format file size in human readable format.
    
    Args:
        bytes_size: Size in bytes
        
    Returns:
        Formatted string like "1.5 MB"
    """
    for unit in ['B', 'KB', 'MB', 'GB']:
        if bytes_size < 1024:
            return f"{bytes_size:.1f} {unit}" if unit != 'B' else f"{bytes_size} {unit}"
        bytes_size /= 1024
    return f"{bytes_size:.1f} TB"


def format_cost(cost_usd: float) -> str:
    """
    Format cost in USD with appropriate precision.
    
    Args:
        cost_usd: Cost in USD
        
    Returns:
        Formatted string
    """
    if cost_usd < 0.01:
        return f"${cost_usd:.6f}"
    elif cost_usd < 1:
        return f"${cost_usd:.4f}"
    else:
        return f"${cost_usd:.2f}"


# =========================================
# Markdown -> Telegram HTML conversion
# =========================================

def convert_markdown_to_html(text: str) -> str:
    """
    Convert Markdown formatting to Telegram HTML.
    Handles: bold, italic, code, code blocks, strikethrough, headers, lists.
    """
    # Escape HTML special characters first
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    
    # Code blocks (```code```) - must be first to avoid conflicts
    text = re.sub(r'```(\w*)\n?(.*?)```', r'<pre>\2</pre>', text, flags=re.DOTALL)
    
    # Inline code (`code`)
    text = re.sub(r'`([^`\n]+)`', r'<code>\1</code>', text)
    
    # Headers (### text) -> bold
    text = re.sub(r'^###\s+(.+)$', r'<b>\1</b>', text, flags=re.MULTILINE)
    text = re.sub(r'^##\s+(.+)$', r'<b>\1</b>', text, flags=re.MULTILINE)
    text = re.sub(r'^#\s+(.+)$', r'<b>\1</b>', text, flags=re.MULTILINE)
    
    # Bold (**text** or __text__)
    text = re.sub(r'\*\*([^*]+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__([^_]+?)__', r'<b>\1</b>', text)
    
    # Italic (*text* or _text_) - careful not to match ** or __
    text = re.sub(r'(?<!\*)\*([^*\n]+?)\*(?!\*)', r'<i>\1</i>', text)
    text = re.sub(r'(?<!_)_([^_\n]+?)_(?!_)', r'<i>\1</i>', text)
    
    # Strikethrough (~~text~~)
    text = re.sub(r'~~([^~]+?)~~', r'<s>\1</s>', text)
    
    return text


# =========================================
# Split long messages for Telegram
# =========================================

TELEGRAM_MSG_LIMIT = 4096
SAFE_MSG_LIMIT = 4000  # Leave some room for HTML tags


def split_text_for_telegram(text: str, limit: int = SAFE_MSG_LIMIT) -> List[str]:
    """
    Split long text into multiple Telegram-friendly chunks.
    Tries to split at paragraph boundaries, then at sentences, then at words.
    
    Args:
        text: Full text to split
        limit: Max chars per message
        
    Returns:
        List of text chunks
    """
    if len(text) <= limit:
        return [text]
    
    chunks = []
    remaining = text
    
    while remaining:
        if len(remaining) <= limit:
            chunks.append(remaining)
            break
        
        # Try to find a good split point
        split_at = limit
        
        # 1. Try to split at paragraph boundary (\n\n)
        para_idx = remaining.rfind('\n\n', 0, limit)
        if para_idx > limit * 0.3:  # At least 30% of the way through
            split_at = para_idx + 2  # Include the newlines
        else:
            # 2. Try to split at newline
            nl_idx = remaining.rfind('\n', 0, limit)
            if nl_idx > limit * 0.3:
                split_at = nl_idx + 1
            else:
                # 3. Try to split at sentence end
                for sep in ['. ', '! ', '? ', '.) ']:
                    sent_idx = remaining.rfind(sep, 0, limit)
                    if sent_idx > limit * 0.3:
                        split_at = sent_idx + len(sep)
                        break
                else:
                    # 4. Split at word boundary
                    space_idx = remaining.rfind(' ', 0, limit)
                    if space_idx > limit * 0.3:
                        split_at = space_idx + 1
                    # else: hard cut at limit
        
        chunks.append(remaining[:split_at].rstrip())
        remaining = remaining[split_at:].lstrip()
    
    return chunks


async def send_long_message(
    message: Message,
    text: str,
    parse_mode: str = "HTML",
    reply_markup=None
) -> List[Message]:
    """
    Send a long text split into multiple Telegram messages.
    HTML conversion is applied. If HTML fails, falls back to plain text.
    reply_markup is attached only to the LAST message.
    
    Returns:
        List of sent Message objects
    """
    # Convert markdown to HTML
    html_text = convert_markdown_to_html(text)
    
    # Split into chunks
    chunks = split_text_for_telegram(html_text)
    
    sent_messages = []
    for i, chunk in enumerate(chunks):
        is_last = (i == len(chunks) - 1)
        markup = reply_markup if is_last else None
        
        try:
            msg = await message.answer(
                chunk,
                parse_mode=parse_mode,
                reply_markup=markup
            )
        except Exception:
            # HTML parse failed — try plain text (strip tags)
            plain = re.sub(r'<[^>]+>', '', chunk)
            msg = await message.answer(
                plain,
                reply_markup=markup
            )
        sent_messages.append(msg)
    
    return sent_messages


async def edit_or_send_long(
    thinking_message: Message,
    original_message: Message,
    text: str,
    parse_mode: str = "HTML",
    reply_markup=None
) -> List[Message]:
    """
    Edit the thinking_message with first chunk of text,
    send remaining chunks as new messages.
    
    Returns:
        List of all messages (edited + new)
    """
    html_text = convert_markdown_to_html(text)
    chunks = split_text_for_telegram(html_text)
    
    result_messages = []
    
    for i, chunk in enumerate(chunks):
        is_last = (i == len(chunks) - 1)
        markup = reply_markup if is_last else None
        
        if i == 0:
            # Edit the thinking message
            try:
                await thinking_message.edit_text(
                    chunk,
                    parse_mode=parse_mode,
                    reply_markup=markup
                )
                result_messages.append(thinking_message)
            except Exception:
                try:
                    plain = re.sub(r'<[^>]+>', '', chunk)
                    await thinking_message.edit_text(plain, reply_markup=markup)
                    result_messages.append(thinking_message)
                except Exception:
                    msg = await original_message.answer(chunk, parse_mode=parse_mode, reply_markup=markup)
                    result_messages.append(msg)
        else:
            # Send new message for continuation
            try:
                msg = await original_message.answer(
                    chunk,
                    parse_mode=parse_mode,
                    reply_markup=markup
                )
            except Exception:
                plain = re.sub(r'<[^>]+>', '', chunk)
                msg = await original_message.answer(plain, reply_markup=markup)
            result_messages.append(msg)
    
    return result_messages


async def send_as_file(
    message: Message,
    text: str,
    filename: str = "response.txt",
    caption: str = None
) -> Message:
    """
    Send text as a downloadable file.
    
    Args:
        message: Message to reply to
        text: Content to put in file
        filename: Name of the file
        caption: Optional caption
        
    Returns:
        Sent message
    """
    file_bytes = text.encode('utf-8')
    doc = BufferedInputFile(file_bytes, filename=filename)
    return await message.answer_document(doc, caption=caption)


def _markdown_to_docx_bytes(text: str) -> bytes:
    """
    Convert markdown/AI-response text to a beautifully styled Word (.docx) document.
    
    Handles: headers (#, ##, ###), bold (**), italic (*), code blocks (```),
    inline code (`), bullet lists (- / *), numbered lists, tables, horizontal rules,
    source links from web search, and proper paragraph spacing.
    
    Returns:
        bytes of the .docx file
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io
    
    doc = Document()
    
    # --- Document-wide styling ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # --- Pre-process: extract source links and clean text ---
    # Extract <a href="url">title</a> links before stripping HTML
    source_links = []
    link_pattern = re.compile(r'<a\s+href="([^"]+)"[^>]*>([^<]+)</a>')
    for m in link_pattern.finditer(text):
        source_links.append({"url": m.group(1), "title": m.group(2)})
    
    # Strip all HTML tags
    clean_text = re.sub(r'<[^>]+>', '', text)
    # Unescape HTML entities
    clean_text = (clean_text.replace('&amp;', '&').replace('&lt;', '<')
                  .replace('&gt;', '>').replace('&#x27;', "'").replace('&quot;', '"'))
    
    lines = clean_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Skip source-link-only lines (🔗 Source | Source)
        if stripped.startswith('🔗') and '|' in stripped:
            i += 1
            continue
        
        # --- Code block ---
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Light gray background
            shading = p.paragraph_format.element.get_or_add_pPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'F5F5F5',
                qn('w:val'): 'clear',
            })
            shading.append(shd)
            continue
        
        # --- Table detection (| col | col |) ---
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:]+\|', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.strip('|').split('|')]
                table_lines.append(cells)
                i += 1
            
            if table_lines:
                cols = max(len(row) for row in table_lines)
                table = doc.add_table(rows=len(table_lines), cols=cols)
                table.style = 'Table Grid'
                
                for r_idx, row_data in enumerate(table_lines):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(2)
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                                    if r_idx == 0:
                                        run.bold = True
                                        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                            # Header row background
                            if r_idx == 0:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = tcPr.makeelement(qn('w:shd'), {
                                    qn('w:fill'): 'E8EAF6',
                                    qn('w:val'): 'clear',
                                })
                                tcPr.append(shd)
            continue
        
        # --- Headers ---
        if stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=3)
            h.paragraph_format.space_before = Pt(8)
            i += 1
            continue
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=2)
            h.paragraph_format.space_before = Pt(12)
            i += 1
            continue
        elif stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=1)
            h.paragraph_format.space_before = Pt(14)
            i += 1
            continue
        
        # --- Bullet list (- item or * item) ---
        if re.match(r'^[\-\*]\s+', stripped):
            item_text = re.sub(r'^[\-\*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            _add_formatted_runs(p, item_text)
            i += 1
            continue
        
        # --- Numbered list (1. item) ---
        if re.match(r'^\d+\.\s+', stripped):
            item_text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            _add_formatted_runs(p, item_text)
            i += 1
            continue
        
        # --- Horizontal rule (--- or ===) ---
        if re.match(r'^[\-]{3,}$', stripped) or re.match(r'^[=]{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Draw a thin line using bottom border
            pPr = p.paragraph_format.element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # --- Regular paragraph ---
        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)
        i += 1
    
    # --- Append source links at the end ---
    if source_links:
        doc.add_paragraph()  # spacer
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run('Источники / Sources')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        for link in source_links:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"• {link['title']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x6B, 0xB8)
            run = p.add_run(f"  ({link['url']})")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_formatted_runs(paragraph, text: str):
    """Add runs to a paragraph with inline markdown formatting (bold, italic, code)."""
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`|\*([^*]+?)\*|_([^_]+?)_)')
    
    last_end = 0
    for match in pattern.finditer(text):
        # Add text before this match as a normal run
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # `code`
            from docx.shared import Pt, RGBColor
            run = paragraph.add_run(match.group(3))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # _italic_
            run = paragraph.add_run(match.group(5))
            run.italic = True
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


async def send_as_docx(
    message: Message,
    text: str,
    filename: str = "response.docx",
    caption: str = None
) -> Message:
    """
    Send text as a nicely formatted Word (.docx) document.
    Converts markdown formatting to Word styles (headers, bold, italic,
    code blocks, tables, bullet/numbered lists).
    
    Args:
        message: Message to reply to
        text: Content (markdown) to format into a Word document
        filename: Name of the .docx file
        caption: Optional caption
        
    Returns:
        Sent message
    """
    try:
        docx_bytes = _markdown_to_docx_bytes(text)
        doc = BufferedInputFile(docx_bytes, filename=filename)
        return await message.answer_document(doc, caption=caption)
    except Exception:
        # Fallback to plain txt if docx generation fails
        file_bytes = text.encode('utf-8')
        fallback_filename = filename.replace('.docx', '.txt')
        doc = BufferedInputFile(file_bytes, filename=fallback_filename)
        return await message.answer_document(doc, caption=caption)
